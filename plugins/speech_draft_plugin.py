import os
import sys
import json

# 确保应用根目录下的 lib 目录在搜索路径中，防止打包后找不到第三方库
def _ensure_lib_path():
    # 1. 尝试获取插件自身的私有 lib 目录 (如果插件以文件夹形式存在)
    current_file_dir = os.path.dirname(os.path.abspath(__file__))
    private_lib = os.path.join(current_file_dir, "lib")
    if os.path.exists(private_lib) and private_lib not in sys.path:
        sys.path.insert(0, private_lib)

    # 2. 尝试获取主应用的 lib 目录
    # 插件文件可能在 plugins/ 或 plugins/folder/ 目录下
    base_dir = current_file_dir
    for _ in range(3): # 最多向上找3级
        if os.path.exists(os.path.join(base_dir, "lib")):
            lib_path = os.path.join(base_dir, "lib")
            if lib_path not in sys.path:
                sys.path.insert(0, lib_path)
            break
        base_dir = os.path.dirname(base_dir)
    
    # 3. 针对 macOS 打包环境 (.app) 的特殊处理
    if getattr(sys, 'frozen', False):
        bundle_dir = os.path.dirname(sys.executable)
        app_lib = os.path.join(bundle_dir, "..", "Resources", "lib")
        if os.path.exists(app_lib) and app_lib not in sys.path:
            sys.path.insert(0, app_lib)

_ensure_lib_path()
# 显式导入 email 及其子模块，防止某些打包环境剔除标准库
try:
    import email
    import email.mime.text
    import email.mime.multipart
except ImportError:
    pass

try:
    import requests
except ImportError:
    requests = None
from src.utils.logger import logger
import threading
from typing import List, Dict, Any, Optional

from PyQt6.QtWidgets import (
    QVBoxLayout, QHBoxLayout, QPushButton, QLineEdit, 
    QTextEdit, QLabel, QMessageBox, QGroupBox, QCheckBox,
    QScrollArea, QWidget, QFileDialog, QProgressBar, QFrame
)
from PyQt6.QtCore import Qt, pyqtSignal, QObject, QSize
from PyQt6.QtGui import QFont, QIcon, QColor

# 尝试导入 python-docx
try:
    from docx import Document
except ImportError:
    Document = None

from src.plugins.base_plugin import BasePlugin

# 尝试导入 openai，如果失败则提示安装
try:
    from openai import OpenAI
except ImportError:
    OpenAI = None


class StreamWorkerSignals(QObject):
    """处理信号"""
    chunk_received = pyqtSignal(str)
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    started = pyqtSignal()
    # 新增：用于单次调用的 UI 更新信号 (内容, 目标控件对象)
    oneshot_finished = pyqtSignal(str, object)

class SpeechDraftPlugin(BasePlugin):
    def __init__(self):
        super().__init__("讲话稿生成", "基于 AI 的讲话稿流式生成工具")
        self.reference_files = []
        self.config_file = self._get_config_path()
        self._setup_ui()
        self._load_config()
        
    def _get_config_path(self):
        """获取配置文件路径"""
        from src.utils.path_utils import get_data_directory
        data_dir = get_data_directory()
        config_dir = os.path.join(data_dir, "speech_draft")
        if not os.path.exists(config_dir):
            os.makedirs(config_dir, exist_ok=True)
        return os.path.join(config_dir, "ai_config.json")

    def _load_config(self):
        """加载配置"""
        if os.path.exists(self.config_file):
            try:
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    config = json.load(f)
                    self.api_base_input.setText(config.get("api_base", "https://api.openai.com/v1"))
                    self.api_key_input.setText(config.get("api_key", ""))
                    self.model_input.setText(config.get("model", "gpt-3.5-turbo"))
                    self.temp_input.setText(str(config.get("temperature", "0.7")))
                    self.max_tokens_input.setText(str(config.get("max_tokens", "2048")))
                    
                    # 加载提示词配置
                    prompts = config.get("prompts", {})
                    kw_p = prompts.get("keywords", {})
                    self.kw_sys_input.setPlainText(kw_p.get("system", "你是一个专业的文案助手。"))
                    self.kw_user_input.setPlainText(kw_p.get("user", "请根据文章标题“{title}”，生成3-5个相关的讲话稿关键词，以逗号分隔。仅返回关键词本身。"))
                    
                    sum_p = prompts.get("summary", {})
                    self.sum_sys_input.setPlainText(sum_p.get("system", "你是一个专业的文案助手。"))
                    self.sum_user_input.setPlainText(sum_p.get("user", "请根据文章标题“{title}”，撰写一段简短的内容概述，描述讲话的场景、听众和核心要点。字数在100字左右。"))
                    
                    gen_p = prompts.get("generation", {})
                    self.gen_sys_input.setPlainText(gen_p.get("system", "你是一个专业的讲话稿写作专家。请根据用户提供的标题、关键词和内容概述，参考提供的联网搜索结果和参考文档，撰写一份高质量、得体、富有感染力的讲话稿。"))
                    self.gen_user_input.setPlainText(gen_p.get("user", "标题：{title}\n关键词：{keywords}\n内容概述：{content}\n联网搜索参考：\n{search_results}\n参考文件：{files}\n\n请开始撰写讲话稿全文："))
                    
                    # 联网搜索配置
                    self.search_url_input.setText(config.get("search_url", "https://open.feedcoopapi.com/search_api/web_search"))
                    self.search_key_input.setText(config.get("search_key", ""))
            except Exception as e:
                logger.error(f"加载 AI 配置失败: {e}")

    def _save_config(self):
        """保存配置"""
        config = {
            "api_base": self.api_base_input.text().strip(),
            "api_key": self.api_key_input.text().strip(),
            "model": self.model_input.text().strip(),
            "temperature": self.temp_input.text().strip(),
            "max_tokens": self.max_tokens_input.text().strip(),
            "search_url": self.search_url_input.text().strip(),
            "search_key": self.search_key_input.text().strip(),
            "prompts": {
                "keywords": {
                    "system": self.kw_sys_input.toPlainText(),
                    "user": self.kw_user_input.toPlainText()
                },
                "summary": {
                    "system": self.sum_sys_input.toPlainText(),
                    "user": self.sum_user_input.toPlainText()
                },
                "generation": {
                    "system": self.gen_sys_input.toPlainText(),
                    "user": self.gen_user_input.toPlainText()
                }
            }
        }
        try:
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(config, f, ensure_ascii=False, indent=4)
        except Exception as e:
            logger.error(f"保存 AI 配置失败: {e}")

    def _setup_ui(self):
        # 初始化信号对象
        self.one_shot_signals = StreamWorkerSignals()
        self.one_shot_signals.oneshot_finished.connect(self._on_oneshot_finished)
        
        # 创建外层布局以包含滚动区域
        outer_layout = QVBoxLayout(self)
        outer_layout.setContentsMargins(0, 0, 0, 0)
        
        # 创建滚动区域
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setFrameShape(QFrame.Shape.NoFrame)
        self.scroll_area.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        
        # 创建内容容器
        container = QWidget()
        container.setObjectName("speech_draft_container")
        
        # 主布局
        main_layout = QVBoxLayout(container)
        main_layout.setContentsMargins(30, 20, 30, 20)
        main_layout.setSpacing(15)
        
        # 风格设置
        self.setStyleSheet("""
            QWidget {
                background-color: #f8f9fa;
                color: #2c3e50;
            }
            QLabel {
                font-weight: bold;
                font-size: 14px;
                color: #34495e;
            }
            QLineEdit, QTextEdit {
                border: 1px solid #dcdfe6;
                border-radius: 6px;
                padding: 8px;
                background-color: #ffffff;
            }
            QLineEdit:focus, QTextEdit:focus {
                border-color: #409eff;
            }
            QPushButton {
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: 500;
            }
            QPushButton#primary_btn {
                background-color: #409eff;
                color: white;
                border: none;
                font-size: 15px;
            }
            QPushButton#primary_btn:hover {
                background-color: #66b1ff;
            }
            QPushButton#secondary_btn {
                background-color: #ffffff;
                border: 1px solid #dcdfe6;
                color: #606266;
            }
            QPushButton#secondary_btn:hover {
                color: #409eff;
                border-color: #c6e2ff;
                background-color: #ecf5ff;
            }
            QPushButton#link_btn {
                background-color: transparent;
                color: #409eff;
                border: none;
                padding: 0;
                font-weight: normal;
                text-decoration: underline;
            }
            QGroupBox {
                border: 1px solid #ebeef5;
                border-radius: 8px;
                margin-top: 10px;
                padding-top: 20px;
                background-color: #ffffff;
            }
            QCheckBox::indicator {
                width: 18px;
                height: 18px;
            }
        """)

        # 1. 文章标题
        title_layout = QVBoxLayout()
        title_header = QHBoxLayout()
        title_label = QLabel("文章标题*")
        self.title_count_label = QLabel("0 / 50")
        self.title_count_label.setStyleSheet("color: #909399; font-weight: normal; font-size: 12px;")
        title_header.addWidget(title_label)
        title_header.addStretch()
        title_header.addWidget(self.title_count_label)
        
        self.title_input = QLineEdit()
        self.title_input.setPlaceholderText("请输入标题，如“在新员工入职欢迎会上的讲话”")
        self.title_input.setMaxLength(50)
        self.title_input.textChanged.connect(self._update_title_count)
        
        title_layout.addLayout(title_header)
        title_layout.addWidget(self.title_input)
        main_layout.addLayout(title_layout)

        # 2. 关键词
        kw_layout = QVBoxLayout()
        kw_header = QHBoxLayout()
        kw_header.addWidget(QLabel("关键词"))
        kw_header.addStretch()
        self.ai_kw_btn = QPushButton("AI写作")
        self.ai_kw_btn.setObjectName("link_btn")
        self.ai_kw_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.ai_kw_btn.clicked.connect(self._ai_generate_keywords)
        kw_header.addWidget(self.ai_kw_btn)
        
        kw_input_layout = QHBoxLayout()
        self.kw_input = QLineEdit()
        self.kw_input.setPlaceholderText("+ 添加关键词 (以逗号或空格分隔)")
        kw_input_layout.addWidget(self.kw_input)
        
        kw_layout.addLayout(kw_header)
        kw_layout.addLayout(kw_input_layout)
        main_layout.addLayout(kw_layout)

        # 3. 内容概述
        content_layout = QVBoxLayout()
        content_header = QHBoxLayout()
        content_header.addWidget(QLabel("内容"))
        
        self.ai_content_btn = QPushButton("AI写内容概述")
        self.ai_content_btn.setObjectName("link_btn")
        self.ai_content_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.ai_content_btn.clicked.connect(self._ai_generate_summary)
        content_header.addWidget(self.ai_content_btn)
        
        content_header.addStretch()
        
        self.content_count_label = QLabel("0 / 2000")
        self.content_count_label.setStyleSheet("color: #909399; font-weight: normal; font-size: 12px;")
        content_header.addWidget(self.content_count_label)
        
        self.content_input = QTextEdit()
        self.content_input.setPlaceholderText("简要填写个人讲话场景、听众，说明讲话核心内容（汇报/表态/分享等）与表达风格")
        self.content_input.setMaximumHeight(120)
        self.content_input.textChanged.connect(self._update_content_count)
        
        content_layout.addLayout(content_header)
        content_layout.addWidget(self.content_input)
        main_layout.addLayout(content_layout)

        # 4. 参考文档
        ref_layout = QVBoxLayout()
        ref_header = QHBoxLayout()
        ref_header.addWidget(QLabel("参考文档"))
        ref_header.addStretch()
        ref_layout.addLayout(ref_header)
        
        self.upload_btn = QPushButton("+ 上传参考文档")
        self.upload_btn.setObjectName("secondary_btn")
        self.upload_btn.setFixedWidth(150)
        self.upload_btn.clicked.connect(self._upload_files)
        
        self.file_list_label = QLabel("未选择文件")
        self.file_list_label.setStyleSheet("color: #909399; font-weight: normal; font-size: 12px;")
        
        ref_layout.addWidget(self.upload_btn)
        ref_layout.addWidget(self.file_list_label)
        main_layout.addLayout(ref_layout)

        # 5. 高级选项
        adv_layout = QHBoxLayout()
        
        self.search_check = QCheckBox("联网搜索")
        self.search_check.setToolTip("开启后会提高文章的时效性和专业程度")
        
        adv_layout.addWidget(self.search_check)
        adv_layout.addStretch()
        main_layout.addLayout(adv_layout)

        # 6. 配置选项 (整合为高级配置)
        self.settings_group = QGroupBox("高级配置")
        self.settings_group.setMinimumHeight(480) # 调大配置区域高度
        settings_main_layout = QVBoxLayout()
        
        from PyQt6.QtWidgets import QTabWidget
        self.config_tabs = QTabWidget()
        
        # --- Tab 1: 模型设置 ---
        model_tab = QWidget()
        model_layout = QVBoxLayout(model_tab)
        
        row1_layout = QHBoxLayout()
        api_base_layout = QVBoxLayout()
        api_base_layout.addWidget(QLabel("API Base"))
        self.api_base_input = QLineEdit("https://ark.cn-beijing.volces.com/api/v3/")
        self.api_base_input.textChanged.connect(self._save_config)
        api_base_layout.addWidget(self.api_base_input)
        
        api_key_layout = QVBoxLayout()
        api_key_layout.addWidget(QLabel("API Key"))
        self.api_key_input = QLineEdit()
        self.api_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.api_key_input.setPlaceholderText("api key...")
        self.api_key_input.textChanged.connect(self._save_config)
        api_key_layout.addWidget(self.api_key_input)
        
        row1_layout.addLayout(api_base_layout, 2)
        row1_layout.addSpacing(15)
        row1_layout.addLayout(api_key_layout, 3)
        
        row2_layout = QHBoxLayout()
        model_name_layout = QVBoxLayout()
        model_name_layout.addWidget(QLabel("Model"))
        self.model_input = QLineEdit("doubao-seed-1-8-251228")
        self.model_input.textChanged.connect(self._save_config)
        model_name_layout.addWidget(self.model_input)
        
        temp_layout = QVBoxLayout()
        temp_layout.addWidget(QLabel("Temperature"))
        self.temp_input = QLineEdit("0.7")
        self.temp_input.textChanged.connect(self._save_config)
        temp_layout.addWidget(self.temp_input)
        
        max_tokens_layout = QVBoxLayout()
        max_tokens_layout.addWidget(QLabel("Max Tokens"))
        self.max_tokens_input = QLineEdit("4096")
        self.max_tokens_input.textChanged.connect(self._save_config)
        max_tokens_layout.addWidget(self.max_tokens_input)
        
        row2_layout.addLayout(model_name_layout, 2)
        row2_layout.addSpacing(15)
        row2_layout.addLayout(temp_layout, 1)
        row2_layout.addSpacing(15)
        row2_layout.addLayout(max_tokens_layout, 1)
        
        model_layout.addLayout(row1_layout)
        model_layout.addSpacing(10)
        model_layout.addLayout(row2_layout)
        
        # --- Tab 2: 提示词配置 ---
        prompt_tab = QScrollArea()
        prompt_tab.setWidgetResizable(True)
        prompt_container = QWidget()
        prompt_layout = QVBoxLayout(prompt_container)
        
        def create_prompt_box(title, sys_attr, user_attr):
            box = QGroupBox(title)
            blayout = QVBoxLayout()
            blayout.addWidget(QLabel("System Prompt:"))
            sys_input = QTextEdit()
            sys_input.setMaximumHeight(60)
            sys_input.textChanged.connect(self._save_config)
            setattr(self, sys_attr, sys_input)
            blayout.addWidget(sys_input)
            
            blayout.addWidget(QLabel("User Prompt Template:"))
            user_input = QTextEdit()
            user_input.setMaximumHeight(100)
            user_input.textChanged.connect(self._save_config)
            setattr(self, user_attr, user_input)
            blayout.addWidget(user_input)
            box.setLayout(blayout)
            return box

        prompt_layout.addWidget(create_prompt_box("关键词生成", "kw_sys_input", "kw_user_input"))
        prompt_layout.addWidget(create_prompt_box("内容概述生成", "sum_sys_input", "sum_user_input"))
        prompt_layout.addWidget(create_prompt_box("讲话稿正文生成", "gen_sys_input", "gen_user_input"))
        prompt_tab.setWidget(prompt_container)
        
        self.config_tabs.addTab(model_tab, "模型设置")
        
        # --- Tab 2: 联网搜索配置 ---
        search_tab = QWidget()
        search_layout = QVBoxLayout(search_tab)
        
        self.search_url_input = QLineEdit("https://open.feedcoopapi.com/search_api/web_search")
        self.search_url_input.setPlaceholderText("API URL")
        self.search_url_input.textChanged.connect(self._save_config)
        search_layout.addWidget(QLabel("联网搜索API URL:"))
        search_layout.addWidget(self.search_url_input)
        
        self.search_key_input = QLineEdit()
        self.search_key_input.setEchoMode(QLineEdit.EchoMode.Password)
        self.search_key_input.setPlaceholderText("API Key")
        self.search_key_input.textChanged.connect(self._save_config)
        search_layout.addWidget(QLabel("Search API Key:"))
        search_layout.addWidget(self.search_key_input)
        search_layout.addStretch()
        
        self.config_tabs.addTab(search_tab, "联网搜索")
        
        # --- Tab 3: 提示词配置 ---
        self.config_tabs.addTab(prompt_tab, "提示词配置")
        
        settings_main_layout.addWidget(self.config_tabs)
        self.settings_group.setLayout(settings_main_layout)
        self.settings_group.setVisible(False)
        
        # 切换设置按钮
        self.toggle_settings_btn = QPushButton("高级配置")
        self.toggle_settings_btn.setObjectName("secondary_btn")
        self.toggle_settings_btn.clicked.connect(lambda: self.settings_group.setVisible(not self.settings_group.isVisible()))
        
        main_layout.addWidget(self.settings_group)
        main_layout.addSpacing(5)

        # 7. 生成按钮
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        
        # 高级设置按钮移到这里
        self.generate_btn = QPushButton("开始生成")
        self.generate_btn.setObjectName("primary_btn")
        self.generate_btn.setMinimumSize(120, 40)
        self.generate_btn.clicked.connect(self._start_generation)
        btn_layout.addWidget(self.generate_btn)
        
        btn_layout.addSpacing(10)
        
        # 高级设置按钮移到生成按钮旁边
        btn_layout.addWidget(self.toggle_settings_btn)
        
        main_layout.addLayout(btn_layout)

        # 8. 输出区域
        self.output_area = QTextEdit()
        self.output_area.setPlaceholderText("生成内容将在这里显示...")
        self.output_area.setReadOnly(True)
        self.output_area.setMinimumHeight(300)
        main_layout.addWidget(self.output_area)

        # 9. 输出底部按钮
        output_footer = QHBoxLayout()
        
        self.copy_btn = QPushButton("复制正文")
        self.copy_btn.setObjectName("secondary_btn")
        self.copy_btn.clicked.connect(self._copy_content)
        
        self.export_btn = QPushButton("导出为 Word")
        self.export_btn.setObjectName("secondary_btn")
        self.export_btn.clicked.connect(self._export_to_word)
        
        output_footer.addStretch()
        output_footer.addWidget(self.copy_btn)
        output_footer.addWidget(self.export_btn)
        main_layout.addLayout(output_footer)
        
        # 将内容容器装载到滚动区域，并将滚动区域放入外层布局
        self.scroll_area.setWidget(container)
        outer_layout.addWidget(self.scroll_area)

    def _copy_content(self):
        """复制内容到剪贴板"""
        text = self.output_area.toPlainText().strip()
        if not text:
            return
        
        from PyQt6.QtWidgets import QApplication
        QApplication.clipboard().setText(text)
        QMessageBox.information(self, "成功", "已复制到剪贴板")

    def _export_to_word(self):
        """导出内容为 Word 文档"""
        text = self.output_area.toPlainText().strip()
        if not text:
            QMessageBox.warning(self, "提醒", "当前没有可导出的内容")
            return

        if Document is None:
            QMessageBox.critical(self, "错误", "未安装 python-docx 库，请运行 'pip install -t lib python-docx'")
            return

        file_path, _ = QFileDialog.getSaveFileName(
            self, "导出讲话稿", f"{self.title_input.text().strip() or '未命名讲话稿'}.docx", "Word Documents (*.docx)"
        )

        if file_path:
            try:
                doc = Document()
                title = self.title_input.text().strip()
                if title:
                    doc.add_heading(title, 0)
                
                # 分段写入
                for paragraph in text.split('\n'):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)
                
                doc.save(file_path)
                QMessageBox.information(self, "成功", f"讲话稿已成功导出至：\n{file_path}")
            except Exception as e:
                logger.error(f"导出 Word 失败: {e}")
                QMessageBox.critical(self, "错误", f"导出失败: {str(e)}")

    def _update_title_count(self):
        count = len(self.title_input.text())
        self.title_count_label.setText(f"{count} / 50")

    def _update_content_count(self):
        count = len(self.content_input.toPlainText())
        self.content_count_label.setText(f"{count} / 1000")

    def _upload_files(self):
        files, _ = QFileDialog.getOpenFileNames(self, "选择参考文档", "", "Documents (*.pdf *.docx *.txt *.md)")
        if files:
            self.reference_files = files
            self.file_list_label.setText(f"已上传 {len(files)} 个文件: " + ", ".join([os.path.basename(f) for f in files]))

    def _ai_generate_keywords(self):
        title = self.title_input.text().strip()
        if not title:
            QMessageBox.warning(self, "提醒", "请先输入文章标题")
            return
        
        sys_p = self.kw_sys_input.toPlainText().strip()
        user_p = self.kw_user_input.toPlainText().strip().format(title=title)
        self._call_ai_oneshot(sys_p, user_p, self.kw_input)

    def _ai_generate_summary(self):
        title = self.title_input.text().strip()
        if not title:
            QMessageBox.warning(self, "提醒", "请先输入文章标题")
            return
        
        keywords = self.kw_input.text().strip()
        sys_p = self.sum_sys_input.toPlainText().strip()
        # 传入 title 和 keywords，同时兼容 keyword 拼写
        user_p = self.sum_user_input.toPlainText().strip().format(
            title=title, 
            keywords=keywords or "无",
            keyword=keywords or "无"
        )
        self._call_ai_oneshot(sys_p, user_p, self.content_input)

    def _on_oneshot_finished(self, content, target_widget):
        """安全地在主线程更新 UI"""
        if isinstance(target_widget, QLineEdit):
            target_widget.setText(content)
        elif isinstance(target_widget, QTextEdit):
            target_widget.setPlainText(content)

    def _call_ai_oneshot(self, system_prompt, user_prompt, target_widget):
        """单次非流式调用 AI (用于生成关键词和概述)"""
        api_key = self.api_key_input.text().strip()
        api_base = self.api_base_input.text().strip()
        model = self.model_input.text().strip()
        
        if not api_key:
            QMessageBox.warning(self, "提醒", "请先在高级配置中设置 API Key")
            self.settings_group.setVisible(True)
            return

        if OpenAI is None:
            QMessageBox.critical(self, "错误", "未安装 openai 库，请在应用根目录下运行 'pip install -t lib openai'")
            return

        def task():
            try:
                logger.info(f"AI单次调用提示词: \nSys: {system_prompt}\nUser: {user_prompt[:500]}")
                client = OpenAI(api_key=api_key, base_url=api_base)
                response = client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=float(self.temp_input.text().strip()) if self.temp_input.text().strip() else 0.7,
                    max_tokens=int(self.max_tokens_input.text().strip()) if self.max_tokens_input.text().strip() else 1000,
                    stream=False
                )
                
                content = response.choices[0].message.content.strip()
                # 通过信号发送结果，而不是直接操作控件
                self.one_shot_signals.oneshot_finished.emit(content, target_widget)
            except Exception as e:
                logger.error(f"AI 调用异常: {e}")

        threading.Thread(target=task, daemon=True).start()

    def _start_generation(self):
        title = self.title_input.text().strip()
        if not title:
            QMessageBox.warning(self, "提醒", "请输入文章标题")
            return

        api_key = self.api_key_input.text().strip()
        if not api_key:
            QMessageBox.warning(self, "提醒", "请先在设置中配置 API Key")
            self.settings_group.setVisible(True)
            return

        if OpenAI is None:
            QMessageBox.critical(self, "错误", "未安装 openai 库，请运行 'pip install -t lib openai'")
            return

        # 禁用按钮
        self.generate_btn.setEnabled(False)
        self.generate_btn.setText("正在生成...")
        self.output_area.clear()

        # 启动处理线程
        threading.Thread(target=self._generation_workflow, args=(title,), daemon=True).start()

    def _generation_workflow(self, title):
        """完整的生成工作流：提取意图 -> 联网搜索 -> 生成正文"""
        keywords = self.kw_input.text().strip()
        content_summary = self.content_input.toPlainText().strip()
        search_results_str = "未开启联网搜索"
        
        # 1. 如果开启了联网搜索
        if self.search_check.isChecked():
            self.one_shot_signals.chunk_received.emit("> 正在提炼搜索意图并联网检索...\n")
            search_results_str = self._perform_online_search(title, keywords, content_summary)
            self.one_shot_signals.chunk_received.emit("> 联网搜索完成，正在组织语言撰写全文...\n\n")

        # 2. 准备最终提示词
        files_str = f"已上传 {len(self.reference_files)} 个相关文档" if self.reference_files else "无"
        
        system_prompt = self.gen_sys_input.toPlainText().strip()
        if self.search_check.isChecked():
            system_prompt += " 请重点参考提供的联网搜索实时信息。"

        user_prompt_template = self.gen_user_input.toPlainText().strip()
        user_prompt = user_prompt_template.format(
            title=title,
            keywords=keywords or "无",
            content=content_summary or "无",
            search_results=search_results_str,
            files=files_str
        )

        # 3. 启动流式生成
        self.signals = StreamWorkerSignals()
        self.signals.chunk_received.connect(self._on_chunk_received)
        self.signals.finished.connect(self._on_finished)
        self.signals.error.connect(self._on_error)
        
        self._stream_task(system_prompt, user_prompt)

    def _perform_online_search(self, title, keywords, summary) -> str:
        """执行联网搜索逻辑"""
        search_url = self.search_url_input.text().strip()
        search_key = self.search_key_input.text().strip()
        
        if not search_url or not search_key:
            return "联网搜索失败：未配置搜索 API URL 或 Key"

        try:
            # A. 提炼搜索意图
            client = OpenAI(api_key=self.api_key_input.text().strip(), base_url=self.api_base_input.text().strip())
            intent_prompt = f"请根据讲话稿标题《{title}》、关键词({keywords})和内容概述({summary})，提炼3个用于联网搜索的精准查询词，以换行分隔。仅返回查询词。"
            
            resp = client.chat.completions.create(
                model=self.model_input.text().strip(),
                messages=[{"role": "user", "content": intent_prompt}],
                temperature=0.5,
                stream=False
            )
            intents = [line.strip() for line in resp.choices[0].message.content.split('\n') if line.strip()][:3]
            logger.info(f"提炼出的搜索意图: {intents}")

            # B. 并行/串行执行搜索
            all_formatted_results = []
            for intent in intents:
                res = self._call_volcano_search(intent, search_url, search_key, count=5)
                if res:
                    all_formatted_results.append(f"--- 关于“{intent}”的搜索结果 ---\n{res}")
            
            return "\n\n".join(all_formatted_results) if all_formatted_results else "联网搜索未找到相关结果。"
            
        except Exception as e:
            logger.error(f"联网搜索流程异常: {e}")
            return f"联网搜索发生异常: {str(e)}"

    def _call_volcano_search(self, query, url, api_key, count=5) -> str:
        """调用火山搜索接口 (同步)"""
        if requests is None:
            return "联网搜索失败：未安装 requests 库，请运行 'pip install -t lib requests'"
        headers = {
            'Authorization': f'Bearer {api_key}',
            'Content-Type': 'application/json'
        }
        data = {
            "Query": query,
            "SearchType": "web",
            "NeedSummary": True,
            "Count": count
        }
        try:
            response = requests.post(url, headers=headers, json=data, timeout=30)
            if response.status_code == 200:
                json_res = response.json()
                # 处理火山特定的错误元数据
                if "ResponseMetadata" in json_res and "Error" in json_res["ResponseMetadata"]:
                    err = json_res["ResponseMetadata"]["Error"]
                    if err: return f"API 报错: {err}"

                results = json_res.get("Result", {}).get("WebResults", [])
                if not results: return ""

                formatted = ""
                for idx, r in enumerate(results, start=1):
                    formatted += f"[{idx}] {r.get('Title')}\n摘要: {r.get('Summary')}\n来源: {r.get('SiteName')} ({r.get('PublishTime')})\n\n"
                return formatted.strip()
            else:
                return f"HTTP {response.status_code}"
        except Exception as e:
            return f"请求异常: {str(e)}"

    def _stream_task(self, system_prompt, user_prompt):
        api_key = self.api_key_input.text().strip()
        api_base = self.api_base_input.text().strip()
        model = self.model_input.text().strip()
        
        logger.info(f"AI写讲话稿 - 系统提示词: {system_prompt}")
        logger.info(f"AI写讲话稿 - 用户提示词: {user_prompt}")
        
        try:
            client = OpenAI(api_key=api_key, base_url=api_base)
            
            # 参考 demo，增加 thinking 禁用配置和更长的 max_tokens 支持
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=float(self.temp_input.text().strip()) if self.temp_input.text().strip() else 0.6,
                max_tokens=int(self.max_tokens_input.text().strip()) if self.max_tokens_input.text().strip() else 4096,
                stream=True,
                extra_body={
                    "thinking": {
                        "type": "disabled"
                    }
                }
            )
            
            for chunk in resp:
                if chunk.choices and chunk.choices[0].delta.content:
                    content = chunk.choices[0].delta.content
                    self.signals.chunk_received.emit(content)
            
            self.signals.finished.emit("生成完成")
            
        except Exception as e:
            self.signals.error.emit(f"AI 调用失败: {str(e)}")

    def _on_chunk_received(self, chunk):
        self.output_area.insertPlainText(chunk)
        # 自动滚动到底部
        self.output_area.verticalScrollBar().setValue(self.output_area.verticalScrollBar().maximum())

    def _on_finished(self, msg):
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("重新生成")

    def _on_error(self, err_msg):
        self.generate_btn.setEnabled(True)
        self.generate_btn.setText("开始生成")
        QMessageBox.critical(self, "错误", f"生成失败：{err_msg}")

    def get_widget(self) -> "SpeechDraftPlugin":
        return self

    def on_activate(self):
        logger.info("讲话稿插件被激活")
        pass

    def on_deactivate(self):
        pass
